#!/usr/bin/env python3
# filepath: /home/nico-kuehn-dci/Desktop/lectures/first_ai/first_ai.py

# Check Python compatibility first
import sys
import warnings

try:
    # Try to import our compatibility layer
    from compat_layer import warn_about_compatibility
    # Warn about any compatibility issues
    is_compatible = warn_about_compatibility()
    if not is_compatible:
        warnings.warn("Critical compatibility issues detected. Some features may not work.")
except ImportError:
    warnings.warn("Compatibility layer not available. Proceeding without compatibility checks.")

# Standard library imports
import re
import os
import datetime
import time
import json
from pathlib import Path

# Try to import groq - handle import errors for different Python versions
try:
    import groq
except ImportError:
    warnings.warn("Failed to import groq module. Attempting alternative imports...")
    try:
        # Try to import from a different location
        import site
        site_packages = site.getsitepackages()[0]
        sys.path.append(site_packages)
        import groq
    except ImportError:
        print("ERROR: Could not import groq module. API functionality will not work.")
        print("Try installing groq with: pip install --no-build-isolation groq==0.24.0")
        groq = None

# Other dependencies
try:
    from dotenv import load_dotenv
    from fpdf import FPDF
    from docx import Document
    from rich.console import Console
    from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TimeElapsedColumn
    from rich.panel import Panel
    from rich.table import Table
    from rich import box
    from rich.markdown import Markdown
except ImportError as e:
    print(f"Error importing dependency: {e}")
    print("Some features may not be available.")

# Import API configuration
try:
    from api_config import get_api_key
except ImportError:
    get_api_key = None
    
# Import path configuration
try:
    from path_config import get_directories
except ImportError:
    get_directories = None
# Import topic manager for system prompts
try:
    from topic_manager import change_topic
    from prompt_manager import get_current_prompt_type, set_current_prompt_type
except ImportError:
    pass
# Import voice interaction capabilities
try:
    import pyttsx3
    import speech_recognition as sr
    import pyaudio  # Required by speech_recognition for microphone input
    VOICE_AVAILABLE = True
    print("Voice capabilities are enabled!")
    
    # Import cloud TTS capabilities
    try:
        import requests
        from playsound import playsound
        import tempfile
        CLOUD_TTS_AVAILABLE = True
        print("Cloud TTS capabilities are enabled!")
    except ImportError as e:
        print(f"Cloud TTS not available: {e}")
        CLOUD_TTS_AVAILABLE = False
    
    # TTS settings
    AUTO_TTS_ENABLED = False  # Only use TTS when user explicitly requests it
    USE_CLOUD_TTS = True      # Prefer cloud TTS when available
    TTS_RATE = 150            # Speech rate (150 is a good default)
    TTS_VOLUME = 1.0          # Volume (0.0 to 1.0)
        
except ImportError as e:
    print(f"Voice capabilities not available: {e}")
    VOICE_AVAILABLE = False
    CLOUD_TTS_AVAILABLE = False
    AUTO_TTS_ENABLED = False
# Import music notation module
try:
    from music_notation import render_abc_notation, extract_abc_notation, get_abc_example
    MUSIC_NOTATION_AVAILABLE = True
except ImportError:
    MUSIC_NOTATION_AVAILABLE = False

# Function to check system audio capabilities
def check_audio_system():
    """Check the system's audio capabilities and report status."""
    status = {
        "tts_basic": False,
        "tts_cloud": False,
        "speech_recognition": False,
        "microphone": False,
        "best_voice_id": None
    }
    
    # Check basic TTS
    try:
        import pyttsx3
        engine = pyttsx3.init()
        status["tts_basic"] = True
        
        # Find the best voice
        voices = engine.getProperty('voices')
        if voices:
            # Prefer female voices as they often sound clearer
            for voice in voices:
                if "female" in voice.name.lower():
                    status["best_voice_id"] = voice.id
                    break
            
            # If no female voice found, use the first voice
            if not status["best_voice_id"] and voices:
                status["best_voice_id"] = voices[0].id
                
        # Clean up
        del engine
    except Exception:
        status["tts_basic"] = False
    
    # Check cloud TTS
    try:
        import requests
        from playsound import playsound
        import tempfile
        
        # Just check if the modules are available, don't make actual requests
        status["tts_cloud"] = True
    except Exception:
        status["tts_cloud"] = False
    
    # Check speech recognition
    try:
        import speech_recognition as sr
        status["speech_recognition"] = True
        
        # Check for microphone
        try:
            recognizer = sr.Recognizer()
            with sr.Microphone() as source:
                recognizer.adjust_for_ambient_noise(source, duration=0.1)
                status["microphone"] = True
        except Exception:
            status["microphone"] = False
    except Exception:
        status["speech_recognition"] = False
    
    return status

# Function for enhanced text-to-speech
def enhanced_tts(text, use_cloud=True, rate=150, volume=1.0):
    """
    Enhanced text-to-speech function with cloud and fallback options.
    
    Args:
        text: Text to speak
        use_cloud: Whether to attempt cloud TTS first
        rate: Speech rate (words per minute)
        volume: Volume level (0.0 to 1.0)
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not VOICE_AVAILABLE:
        console.print("[red]Voice capabilities not available.[/red]")
        return False
    
    # Clean the text for better speech synthesis
    # Remove markdown formatting, code blocks, etc.
    clean_text = re.sub(r'```.*?```', 'code block', text, flags=re.DOTALL)
    clean_text = re.sub(r'`.*?`', '', clean_text)
    clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
    clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
    
    # Try cloud TTS if requested and available
    if use_cloud and CLOUD_TTS_AVAILABLE:
        console.print("[yellow]Using cloud TTS...[/yellow]")
        
        try:
            # Maximum length per request
            MAX_LENGTH = 200
            # Split text into sentences for more natural breaks
            import re
            sentences = re.split(r'(?<=[.!?])\s+', clean_text)
            
            # Group sentences into chunks under MAX_LENGTH
            chunks = []
            current_chunk = ""
            
            for sentence in sentences:
                if len(current_chunk) + len(sentence) <= MAX_LENGTH:
                    current_chunk += sentence + " "
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + " "
            
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[bold blue]Generating speech..."),
                BarColumn(),
                TimeElapsedColumn()
            ) as progress:
                task = progress.add_task("[green]Processing...", total=len(chunks))
                
                for i, chunk in enumerate(chunks):
                    # Google Translate TTS API - no API key required
                    tts_url = f"https://translate.google.com/translate_tts"
                    params = {
                        "ie": "UTF-8",
                        "q": chunk,
                        "tl": "en",
                        "client": "tw-ob"
                    }
                    
                    response = requests.get(tts_url, params=params)
                    
                    if response.status_code == 200:
                        # Create temp file for the audio
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as f:
                            f.write(response.content)
                            audio_file = f.name
                        
                        # Play the audio
                        playsound(audio_file)
                        
                        # Clean up the temp file
                        try:
                            os.remove(audio_file)
                        except:
                            pass
                    
                    progress.update(task, advance=1)
            
            return True
                
        except Exception as e:
            console.print(f"[bold red]Error with cloud TTS:[/bold red] {e}")
            console.print("[yellow]Falling back to basic TTS...[/yellow]")
    
    # Use basic TTS as fallback
    try:
        console.print("[yellow]Using basic TTS...[/yellow]")
        engine = pyttsx3.init()
        
        # Try to use a better voice if available
        voices = engine.getProperty('voices')
        if len(voices) > 1:  # If multiple voices are available
            # Try to find a female voice which is often clearer
            for voice in voices:
                if "female" in voice.name.lower():
                    engine.setProperty('voice', voice.id)
                    break
        
        # Adjust rate for better clarity
        engine.setProperty('rate', rate)
        engine.setProperty('volume', volume)
        
        engine.say(clean_text)
        engine.runAndWait()
        return True
    
    except Exception as e:
        console.print(f"[bold red]Error with basic TTS:[/bold red] {e}")
        return False

# Function to load system prompts from JSON file
def load_system_prompts():
    try:
        import json
        current_dir = os.path.dirname(os.path.abspath(__file__))
        prompts_file = os.path.join(current_dir, "system_prompts.json")
        
        if os.path.exists(prompts_file):
            with open(prompts_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            return {
                "general": "You are a professional Assistant and Tutor in music theory and composition. You are very helpful and knowledgeable. You are very patient and kind. You are very good at explaining things. Use Markdown formatting in your responses."
            }
    except Exception as e:
        console.print(f"[bold red]Error loading system prompts:[/bold red] {e}")
        return {
            "general": "You are a professional Assistant and Tutor in music theory and composition. You are very helpful and knowledgeable. You are very patient and kind. You are very good at explaining things."
        }

# Load system prompts
SYSTEM_PROMPTS = load_system_prompts()
# Use the prompt manager for tracking the current prompt type
try:
    from prompt_manager import get_current_prompt_type
    # The current prompt type will be handled by the prompt_manager module
except ImportError:
    # Fallback if prompt_manager is not available
    def get_current_prompt_type():
        return "general"
    
    def set_current_prompt_type(prompt_type):
        return prompt_type

# Create console for rich output
console = Console()

# Define available models
MODELS = {
    "1": {"id": "llama-3.3-70b-versatile", "name": "Llama 3.3 70B (Versatile)"},
    "2": {"id": "llama-3.1-70b-versatile", "name": "Llama 3.1 70B (Versatile)"},
    "3": {"id": "mixtral-8x7b-32768", "name": "Mixtral 8x7B-32K"},
    "4": {"id": "gemma-7b-it", "name": "Gemma 7B-IT"}
}

# Default model and temperature
current_model = "llama-3.3-70b-versatile"
temperature = 0.7

load_dotenv()

# Check if API key is available
api_key = os.getenv("GROQ_API_KEY")

# If api_key not found in environment variables, try our custom API key management
if not api_key and get_api_key:
    api_key = get_api_key()

# Final check if we have an API key
if not api_key:
    console.print("[bold red]Error:[/bold red] GROQ_API_KEY not found.")
    console.print("Please set your API key using one of the following methods:")
    console.print("1. Create a .env file with GROQ_API_KEY=your_api_key")
    console.print("2. Set an environment variable: export GROQ_API_KEY=your_api_key")
    console.print("3. Run music-theory-ai-config (if installed as a Debian package)")
    exit(1)

client = groq.Client(api_key=api_key)
# Setting up the conversation
current_topic = get_current_prompt_type()  # Get current topic from prompt manager
conversation = [
    {"role": "system",
     "content": SYSTEM_PROMPTS[current_topic]}
]

# Function to extract and format URLs from text
def extract_urls(text):
    import re
    url_pattern = r'(https?://[^\s]+)'
    return re.findall(url_pattern, text)

# Function to speak text using TTS
def speak_text(text, use_cloud=None, show_progress=True):
    """
    Speak text using text-to-speech.
    
    Args:
        text (str): The text to speak
        use_cloud (bool, optional): Whether to use cloud TTS. If None, uses the default setting.
        show_progress (bool): Whether to show progress indicators
    """
    if not VOICE_AVAILABLE:
        return False
        
    # Determine whether to use cloud TTS
    should_use_cloud = use_cloud if use_cloud is not None else (USE_CLOUD_TTS and CLOUD_TTS_AVAILABLE)
    
    if should_use_cloud and CLOUD_TTS_AVAILABLE:
        # Use cloud-based TTS
        if show_progress:
            console.print("[yellow]Using cloud TTS...[/yellow]")
        
        try:
            # Maximum length per request
            MAX_LENGTH = 200
            chunks = [text[i:i+MAX_LENGTH] for i in range(0, len(text), MAX_LENGTH)]
            
            if show_progress:
                progress_indicator = Progress(
                    SpinnerColumn(),
                    TextColumn("[bold blue]Generating speech..."),
                    BarColumn(),
                    TimeElapsedColumn()
                )
            else:
                # Create a dummy progress context manager that does nothing
                class DummyProgress:
                    def __enter__(self): return self
                    def __exit__(self, *args): pass
                    def add_task(self, *args, **kwargs): return 0
                    def update(self, *args, **kwargs): pass
                progress_indicator = DummyProgress()
            
            with progress_indicator as progress:
                if show_progress:
                    task = progress.add_task("[green]Processing...", total=len(chunks))
                
                for i, chunk in enumerate(chunks):
                    # Google Translate TTS API - no API key required
                    tts_url = f"https://translate.google.com/translate_tts"
                    params = {
                        "ie": "UTF-8",
                        "q": chunk,
                        "tl": "en",
                        "client": "tw-ob"
                    }
                    
                    response = requests.get(tts_url, params=params)
                    
                    if response.status_code == 200:
                        # Create temp file for the audio
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as f:
                            f.write(response.content)
                            audio_file = f.name
                        
                        # Play the audio
                        playsound(audio_file)
                        
                        # Clean up the temp file
                        try:
                            os.remove(audio_file)
                        except:
                            pass
                    
                    if show_progress:
                        progress.update(task, advance=1)
            
            return True
        except Exception as e:
            if show_progress:
                console.print(f"[bold red]Error with cloud TTS:[/bold red] {e}")
                console.print("[yellow]Falling back to basic TTS...[/yellow]")
            # Fall back to basic TTS
            return speak_text(text, use_cloud=False, show_progress=show_progress)
    else:
        # Use basic TTS
        if show_progress:
            console.print("[yellow]Using basic TTS...[/yellow]")
        
        try:
            engine = pyttsx3.init()
            
            # Try to use a better voice if available
            voices = engine.getProperty('voices')
            if len(voices) > 1:  # If multiple voices are available
                # Try to find a female voice which is often clearer
                for voice in voices:
                    if "female" in voice.name.lower():
                        engine.setProperty('voice', voice.id)
                        break
            
            # Adjust rate for better clarity
            engine.setProperty('rate', TTS_RATE)
            engine.setProperty('volume', TTS_VOLUME)
            engine.say(text)
            engine.runAndWait()
            return True
        except Exception as e:
            if show_progress:
                console.print(f"[bold red]Error with basic TTS:[/bold red] {e}")
            return False

# Function to format text with clickable links
def format_with_clickable_links(text):
    import re
    url_pattern = r'(https?://[^\s]+)'
    
    # Replace URLs with rich formatted links
    def replace_with_link(match):
        url = match.group(1)
        return f"[bold blue][link={url}]{url}[/link][/bold blue]"
    
    return re.sub(url_pattern, replace_with_link, text)

# Function to save conversation to text file
def save_to_txt(conversation_history, filename=None):
    try:
        if filename is None:
            # Get the appropriate directory for saved chats
            if get_directories:
                save_dir, _ = get_directories()
            else:
                # Fallback to local directory
                current_dir = os.path.dirname(os.path.abspath(__file__))
                save_dir = os.path.join(current_dir, "saved_chats")
                os.makedirs(save_dir, exist_ok=True)
                
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.join(save_dir, f"music_theory_chat_{timestamp}.txt")
        
        console.print(f"[blue]Attempting to save to file:[/blue] {filename}")
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Saving text file..."),
            BarColumn(),
            TimeElapsedColumn()
        ) as progress:
            task = progress.add_task("[green]Writing...", total=100)
            
            # Start the task
            progress.update(task, advance=10)
            time.sleep(0.2)  # Simulate initial setup
            
            with open(filename, 'w', encoding='utf-8') as file:
                file.write("Music Theory AI Chat\n")
                file.write("="*50 + "\n\n")
                
                progress.update(task, advance=20)
                time.sleep(0.1)  # Simulate processing
                
                # Calculate message count for progress updates
                message_count = len(conversation_history[1:])
                progress_per_message = 60 / max(1, message_count)
                
                # Skip the system message
                for message in conversation_history[1:]:
                    role = message["role"].capitalize()
                    content = message["content"]
                    
                    # Add any URLs as references at the end of messages
                    urls = extract_urls(content)
                    if urls and role.lower() == "assistant":
                        file.write(f"{role}: {content}\n\n")
                        file.write("Links referenced:\n")
                        for i, url in enumerate(urls, 1):
                            file.write(f"{i}. {url}\n")
                        file.write("\n")
                    else:
                        file.write(f"{role}: {content}\n\n")
                        
                    progress.update(task, advance=progress_per_message)
                    time.sleep(0.05)  # Small delay for visual effect
            
            # Finish up
            progress.update(task, advance=10)
            time.sleep(0.2)
            progress.update(task, completed=100)
        
        console.print(f"[bold green]Success:[/bold green] Text file saved to {filename}")
        return filename
    except Exception as e:
        console.print(f"[bold red]Error saving file:[/bold red] {e}")
        return None

# Function to save conversation to PDF
def save_to_pdf(conversation_history, filename=None):
    
    try:
        if filename is None:
            # Get the appropriate directory for saved chats
            if get_directories:
                save_dir, _ = get_directories()
            else:
                # Fallback to local directory
                current_dir = os.path.dirname(os.path.abspath(__file__))
                save_dir = os.path.join(current_dir, "saved_chats")
                os.makedirs(save_dir, exist_ok=True)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.join(save_dir, f"music_theory_chat_{timestamp}.pdf")
        
        console.print(f"[blue]Attempting to save PDF to:[/blue] {filename}")
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Creating PDF..."),
            BarColumn(),
            TimeElapsedColumn()
        ) as progress:
            task = progress.add_task("[green]Processing...", total=100)
            
            # Start the task
            progress.update(task, advance=10)
            time.sleep(0.3)  # PDF initialization takes a bit longer
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            progress.update(task, advance=10)
            
            # Add title
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="Music Theory AI Chat", ln=True, align='C')
            pdf.ln(10)
            
            # Reset font
            pdf.set_font("Arial", size=12)
            
            progress.update(task, advance=10)
            
            # Calculate message count for progress updates
            message_count = len(conversation_history[1:])
            progress_per_message = 60 / max(1, message_count)
            
            # Skip the system message
            for message in conversation_history[1:]:
                role = message["role"].capitalize()
                content = message["content"]
                
                # Add role as bold
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(20, 10, txt=f"{role}:", ln=False)
                
                # Reset font for content
                pdf.set_font("Arial", size=12)
                
                # Handle multi-line content
                pdf.multi_cell(0, 10, txt=content)
                pdf.ln(5)
                
                # Update progress
                progress.update(task, advance=progress_per_message)
                time.sleep(0.1)  # PDF processing is slower
            
            # Finalize PDF
            progress.update(task, advance=5)
            pdf.output(filename)
            
            # Complete
            progress.update(task, advance=5)
            time.sleep(0.2)
            progress.update(task, completed=100)
        
        console.print(f"[bold green]Success:[/bold green] PDF file saved to {filename}")
        return filename
    except Exception as e:
        console.print(f"[bold red]Error saving PDF:[/bold red] {e}")
        return None

# Function to save session for later loading
def save_session(conversation_history, filename=None):
    try:
        if filename is None:
            # Get the appropriate directory for saved sessions
            if get_directories:
                _, save_dir = get_directories()
            else:
                # Fallback to local directory
                current_dir = os.path.dirname(os.path.abspath(__file__))
                save_dir = os.path.join(current_dir, "saved_sessions")
                os.makedirs(save_dir, exist_ok=True)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.join(save_dir, f"session_{timestamp}.json")
        
        console.print(f"[blue]Saving session to:[/blue] {filename}")
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Saving session..."),
            BarColumn(),
            TimeElapsedColumn()
        ) as progress:
            task = progress.add_task("[green]Writing...", total=100)
            
            # Start the task
            progress.update(task, advance=30)
            time.sleep(0.2)  # Simulate setup
            
            import json
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(conversation_history, f, ensure_ascii=False, indent=2)
            
            # Finish up
            progress.update(task, advance=70)
            time.sleep(0.2)
            progress.update(task, completed=100)
        
        console.print(f"[bold green]Success:[/bold green] Session saved to {filename}")
        return filename
    except Exception as e:
        console.print(f"[bold red]Error saving session:[/bold red] {e}")
        return None

# Function to load a previous session
def load_session(filename=None):
    try:
        if filename is None:
            # Get the appropriate directory for saved sessions
            if get_directories:
                _, sessions_dir = get_directories()
            else:
                # Fallback to local directory
                current_dir = os.path.dirname(os.path.abspath(__file__))
                sessions_dir = os.path.join(current_dir, "saved_sessions")
            
            if not os.path.exists(sessions_dir):
                console.print("[yellow]No saved sessions found.[/yellow]")
                return None
                
            sessions = sorted([f for f in os.listdir(sessions_dir) if f.endswith('.json')], 
                            reverse=True)
            
            if not sessions:
                console.print("[yellow]No saved sessions found.[/yellow]")
                return None
            
            # Create a table to display available sessions
            from rich.table import Table
            table = Table(title="Available Sessions")
            table.add_column("Option", style="cyan", justify="right")
            table.add_column("Session Name", style="green")
            
            for i, session in enumerate(sessions, 1):
                # Clean up the filename for display
                display_name = session.replace('session_', '').replace('.json', '')
                display_name = f"{display_name[:8]} {display_name[9:11]}:{display_name[11:13]}"
                table.add_row(str(i), display_name)
            
            console.print(table)
            choice = input("Select a session (or press Enter to cancel): ")
            
            if not choice or not choice.isdigit() or int(choice) < 1 or int(choice) > len(sessions):
                console.print("[yellow]Session loading cancelled.[/yellow]")
                return None
                
            filename = os.path.join(sessions_dir, sessions[int(choice) - 1])
        
        console.print(f"[blue]Loading session from:[/blue] {filename}")
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Loading session..."),
            BarColumn(),
            TimeElapsedColumn()
        ) as progress:
            task = progress.add_task("[green]Reading...", total=100)
            
            progress.update(task, advance=30)
            time.sleep(0.2)  # Simulate reading delay
            
            import json
            with open(filename, 'r', encoding='utf-8') as f:
                loaded_conversation = json.load(f)
            
            progress.update(task, advance=70)
            time.sleep(0.2)
            progress.update(task, completed=100)
        
        console.print(f"[bold green]Success:[/bold green] Session loaded from {filename}")
        return loaded_conversation
    except Exception as e:
        console.print(f"[bold red]Error loading session:[/bold red] {e}")
        return None

# Function to save conversation to DOCX
def save_to_docx(conversation_history, filename=None):
    try:
        if filename is None:
            # Get the appropriate directory for saved chats
            if get_directories:
                save_dir, _ = get_directories()
            else:
                # Fallback to local directory
                current_dir = os.path.dirname(os.path.abspath(__file__))
                save_dir = os.path.join(current_dir, "saved_chats")
                os.makedirs(save_dir, exist_ok=True)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.join(save_dir, f"music_theory_chat_{timestamp}.docx")
        
        console.print(f"[blue]Attempting to save DOCX to:[/blue] {filename}")
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Creating Word document..."),
            BarColumn(),
            TimeElapsedColumn()
        ) as progress:
            task = progress.add_task("[green]Processing...", total=100)
            
            # Start the task
            progress.update(task, advance=15)
            time.sleep(0.2)  # Document initialization
            
            doc = Document()
            doc.add_heading('Music Theory AI Chat', 0)
            
            progress.update(task, advance=15)
            
            # Calculate message count for progress updates
            message_count = len(conversation_history[1:])
            progress_per_message = 60 / max(1, message_count)
            
            # Skip the system message
            for message in conversation_history[1:]:
                role = message["role"].capitalize()
                content = message["content"]
                
                p = doc.add_paragraph()
                p.add_run(f"{role}: ").bold = True
                p.add_run(content)
                
                # Add space between messages
                doc.add_paragraph()
                
                # Update progress
                progress.update(task, advance=progress_per_message)
                time.sleep(0.07)  # Small delay for visual effect
            
            # Save the document
            progress.update(task, advance=10)
            doc.save(filename)
            
            # Complete
            time.sleep(0.2)
            progress.update(task, completed=100)
        
        console.print(f"[bold green]Success:[/bold green] Word document saved to {filename}")
        return filename
    except Exception as e:
        console.print(f"[bold red]Error saving DOCX:[/bold red] {e}")
        return None

from rich.panel import Panel
from rich.table import Table
from rich import box

# Create a welcome message
welcome_message = (
    "[bold green]Welcome to your Music Theory AI Chat![/bold green]\n\n"
    "[yellow]Type your questions about music theory and composition to get started.[/yellow]\n"
    "[yellow]Try the 'topic' command to specialize in areas like harmony or ear training.[/yellow]\n"
    + ("[yellow]Voice commands available! Try 'voice input' and 'voice output'.[/yellow]\n" if VOICE_AVAILABLE else "") +
    "[yellow]Type 'help' to see all available commands.[/yellow]"
)

# Create a welcome panel
welcome_panel = Panel(
    welcome_message,
    title="🎵 Music Theory Assistant 🎹",
    subtitle="Powered by Groq AI",
    border_style="green",
    box=box.ROUNDED
)

console.print(welcome_panel)
console.print("-" * 50)

while True:
    user_message = input("\nYou: ")
    
    # Check for commands
    if user_message.lower() == 'exit':
        console.print("[bold green]Goodbye! Chat ended.[/bold green]")
        break
    elif user_message.lower() == 'help':
        from rich.panel import Panel
        from rich.table import Table
        
        # Create a table for commands
        table = Table(title="Available Commands")
        table.add_column("Command", style="cyan")
        table.add_column("Description", style="green")
        
        # Add command rows
        table.add_row("exit", "End the conversation")
        table.add_row("save txt", "Save conversation as text file")
        table.add_row("save pdf", "Save conversation as PDF file") 
        table.add_row("save docx", "Save conversation as Word document")
        table.add_row("help", "Show this help menu")
        table.add_row("clear", "Clear the conversation history")
        table.add_row("models", "List and select available AI models")
        table.add_row("temp <value>", "Set temperature (0.0-1.0)")
        table.add_row("save session", "Save current conversation to a file")
        table.add_row("load session", "Load a previously saved conversation")
        table.add_row("topic", "Select a specialized music topic")
        if MUSIC_NOTATION_AVAILABLE:
            table.add_row("music example", "Show an example music notation")
            table.add_row("render notation", "Render ABC notation from the last AI response")
        if VOICE_AVAILABLE:
            table.add_row("voice input", "Use voice input for your message")
            table.add_row("voice output", "Read the last AI response aloud with voice options")
            table.add_row("voice settings", "Configure voice output settings")
            table.add_row("install cloud tts", "Install packages required for high-quality voice")
        
        # Display the help panel with the table inside
        console.print(Panel(table, title="Music Theory AI Chat - Help Menu", border_style="blue"))
        continue
    elif user_message.lower() == 'clear':
        # Keep system message but clear conversation history
        conversation = [conversation[0]]
        console.print("[yellow]Conversation history cleared.[/yellow]")
        continue
    elif user_message.lower() == 'models':
        from rich.table import Table
        
        # Create a table for models
        table = Table(title="Available AI Models")
        table.add_column("Option", style="cyan", justify="center")
        table.add_column("Model", style="green")
        table.add_column("Description", style="yellow")
        
        # Add model rows
        for key, model in MODELS.items():
            marker = "→" if model["id"] == current_model else " "
            table.add_row(f"{key} {marker}", model["name"], "Currently selected" if model["id"] == current_model else "")
        
        console.print(table)
        model_choice = input("Select model number (or press Enter to keep current): ")
        if model_choice in MODELS:
            current_model = MODELS[model_choice]["id"]
            console.print(f"[green]Switched to model:[/green] {MODELS[model_choice]['name']}")
        continue
    elif user_message.lower().startswith('temp '):
        try:
            temp_value = float(user_message.split(' ')[1])
            if 0 <= temp_value <= 1:
                temperature = temp_value
                console.print(f"[green]Temperature set to:[/green] {temperature}")
            else:
                console.print("[red]Temperature must be between 0.0 and 1.0[/red]")
        except (IndexError, ValueError):
            console.print("[red]Invalid format. Use 'temp 0.7' (value between 0.0 and 1.0)[/red]")
        continue
    elif user_message.lower() == 'save txt':
        filename = save_to_txt(conversation)
        if filename:
            console.print(f"[green]Conversation saved to[/green] {filename}")
        else:
            console.print("[bold red]Failed to save conversation to text file.[/bold red]")
        continue
    elif user_message.lower() == 'save pdf':
        filename = save_to_pdf(conversation)
        if filename:
            console.print(f"[green]Conversation saved to[/green] {filename}")
        else:
            console.print("[bold red]Failed to save conversation to PDF file.[/bold red]")
        continue
    elif user_message.lower() == 'save docx':
        filename = save_to_docx(conversation)
        if filename:
            console.print(f"[green]Conversation saved to[/green] {filename}")
        else:
            console.print("[bold red]Failed to save conversation to DOCX file.[/bold red]")
        continue
    elif user_message.lower() == 'save session':
        filename = save_session(conversation)
        if filename:
            console.print(f"[green]Session saved to[/green] {filename}")
        else:
            console.print("[bold red]Failed to save session.[/bold red]")
        continue
    elif user_message.lower() == 'load session':
        loaded_conversation = load_session()
        if loaded_conversation:
            conversation = loaded_conversation
            console.print("[green]Session loaded successfully![/green]")
        continue
    elif MUSIC_NOTATION_AVAILABLE and user_message.lower() == 'music example':
        # Show music notation examples
        example_types = ["scale", "chord", "melody"]
        console.print("[yellow]Available example types:[/yellow]")
        for i, ex_type in enumerate(example_types, 1):
            console.print(f"  {i}. {ex_type.capitalize()}")
        
        choice = input("Select an example type (or press Enter for scale): ")
        
        if not choice or not choice.isdigit() or int(choice) < 1 or int(choice) > len(example_types):
            example_type = "scale"
        else:
            example_type = example_types[int(choice) - 1]
        
        # Get the example
        abc_notation = get_abc_example(example_type)
        console.print(Panel(abc_notation, title=f"{example_type.capitalize()} Example (ABC Notation)", border_style="cyan"))
        
        # Render the example
        with Progress(
            SpinnerColumn(),
            TextColumn("[bold blue]Rendering music notation..."),
            BarColumn(),
            TimeElapsedColumn()
        ) as progress:
            task = progress.add_task("[green]Rendering...", total=100)
            
            progress.update(task, advance=30)
            time.sleep(0.3)
            
            image_path = render_abc_notation(abc_notation)
            
            progress.update(task, advance=70)
            time.sleep(0.2)
            progress.update(task, completed=100)
        
        if image_path:
            console.print(f"[green]Music notation rendered to:[/green] {image_path}")
            console.print("[yellow]You can view this image file to see the rendered notation.[/yellow]")
        else:
            console.print("[bold red]Failed to render music notation.[/bold red]")
        continue
    elif MUSIC_NOTATION_AVAILABLE and user_message.lower() == 'render notation':
        # Extract ABC notation from the last AI response
        if len(conversation) < 2 or conversation[-1]["role"] != "assistant":
            console.print("[yellow]No AI response available to extract notation from.[/yellow]")
            continue
        
        last_response = conversation[-1]["content"]
        abc_blocks = extract_abc_notation(last_response)
        
        if not abc_blocks:
            console.print("[yellow]No ABC notation found in the last AI response.[/yellow]")
            console.print("[yellow]Ask the AI to provide music examples in ABC notation format.[/yellow]")
            continue
        
        console.print(f"[green]Found {len(abc_blocks)} ABC notation blocks in the response.[/green]")
        
        # Process each block
        for i, abc_block in enumerate(abc_blocks, 1):
            console.print(f"\n[bold]Block {i}:[/bold]")
            console.print(Panel(abc_block, title=f"ABC Notation Block {i}", border_style="cyan"))
            
            # Render the notation
            with Progress(
                SpinnerColumn(),
                TextColumn(f"[bold blue]Rendering block {i}..."),
                BarColumn(),
                TimeElapsedColumn()
            ) as progress:
                task = progress.add_task("[green]Rendering...", total=100)
                
                progress.update(task, advance=30)
                time.sleep(0.3)
                
                image_path = render_abc_notation(abc_block)
                
                progress.update(task, advance=70)
                time.sleep(0.2)
                progress.update(task, completed=100)
            
            if image_path:
                console.print(f"[green]Music notation rendered to:[/green] {image_path}")
            else:
                console.print("[bold red]Failed to render this ABC notation block.[/bold red]")
        continue
    elif user_message.lower() == 'topic':
        # Use the topic manager to change the current topic
        try:
            current_topic = get_current_prompt_type()
            new_topic, changed = change_topic(conversation, current_topic, SYSTEM_PROMPTS)
            if changed:
                set_current_prompt_type(new_topic)  # Update the current topic using prompt manager
        except Exception as e:
            console.print(f"[bold red]Error changing topic:[/bold red] {e}")
        continue
    elif user_message.lower() == 'install cloud tts':
        console.print("[yellow]Installing required packages for cloud TTS...[/yellow]")
        
        try:
            import subprocess
            result = subprocess.run(['pip', 'install', 'requests', 'playsound'], 
                                   capture_output=True, text=True)
                                   
            if result.returncode == 0:
                console.print("[bold green]Successfully installed cloud TTS packages![/bold green]")
                console.print("[yellow]Please restart the application to use cloud TTS.[/yellow]")
            else:
                console.print("[bold red]Error installing packages:[/bold red]")
                console.print(result.stderr)
        except Exception as e:
            console.print(f"[bold red]Installation error:[/bold red] {e}")
        continue
    elif user_message.lower() == 'voice settings':
        if not VOICE_AVAILABLE:
            console.print("[red]Voice capabilities not available.[/red]")
            continue
        
        console.print("[yellow]Voice Settings:[/yellow]")
        
        engine = pyttsx3.init()
        current_rate = engine.getProperty('rate')
        current_volume = engine.getProperty('volume')
        
        # Display available voices
        voices = engine.getProperty('voices')
        console.print("[cyan]Available Voices:[/cyan]")
        for i, voice in enumerate(voices):
            console.print(f"{i+1}. {voice.name}")
        
        # Voice selection
        voice_choice = input("Select voice number (or Enter to skip): ")
        if voice_choice.isdigit() and 1 <= int(voice_choice) <= len(voices):
            engine.setProperty('voice', voices[int(voice_choice)-1].id)
        
        # Rate adjustment
        console.print(f"[cyan]Current Rate:[/cyan] {current_rate} (normal is ~200, lower is slower)")
        rate_choice = input("Enter new rate (or Enter to skip): ")
        if rate_choice.isdigit():
            engine.setProperty('rate', int(rate_choice))
        
        # Volume adjustment
        console.print(f"[cyan]Current Volume:[/cyan] {current_volume} (0.0-1.0)")
        vol_choice = input("Enter new volume (or Enter to skip): ")
        try:
            vol = float(vol_choice)
            if 0.0 <= vol <= 1.0:
                engine.setProperty('volume', vol)
        except:
            pass
        
        # Test the voice
        console.print("[green]Testing voice with new settings...[/green]")
        engine.say("This is a test of the music theory AI voice with the current settings.")
        engine.runAndWait()
        continue
    elif VOICE_AVAILABLE and user_message.lower() == 'voice input':
        console.print("[yellow]Listening... (speak now)[/yellow]")
        recognizer = sr.Recognizer()
        
        # Set up retry mechanism
        max_retries = 2
        retry_count = 0
        
        while retry_count <= max_retries:
            try:
                with sr.Microphone() as source:
                    # Adjust recognition parameters for better results
                    recognizer.adjust_for_ambient_noise(source, duration=0.5)
                    recognizer.energy_threshold = 300  # Increase energy threshold for better detection
                    recognizer.dynamic_energy_threshold = True
                    
                    console.print("[yellow]Listening... (speak now)[/yellow]")
                    audio = recognizer.listen(source, timeout=5, phrase_time_limit=10)
                    
                    console.print("[yellow]Processing speech...[/yellow]")
                    # Try Google's service first as it's generally more accurate
                    try:
                        speech_text = recognizer.recognize_google(audio)
                    except:
                        # Fall back to local recognition if Google service fails
                        speech_text = recognizer.recognize_sphinx(audio) if hasattr(recognizer, 'recognize_sphinx') else ""
                        
                    if speech_text:
                        console.print(f"[green]You said:[/green] {speech_text}")
                        # Use the recognized text as the user message
                        user_message = speech_text
                        # Continue with normal message processing (don't use continue here)
                        break
                    else:
                        raise sr.UnknownValueError("Empty recognition result")
                        
            except sr.WaitTimeoutError:
                retry_count += 1
                if retry_count <= max_retries:
                    console.print(f"[yellow]No speech detected. Retry {retry_count}/{max_retries}...[/yellow]")
                else:
                    console.print("[red]No speech detected after multiple tries. Please try again later.[/red]")
                    continue
            except sr.UnknownValueError:
                retry_count += 1
                if retry_count <= max_retries:
                    console.print(f"[yellow]Could not understand audio. Retry {retry_count}/{max_retries}...[/yellow]")
                else:
                    console.print("[red]Could not understand audio after multiple tries.[/red]")
                    console.print("[yellow]Tip: Speak clearly and ensure your microphone is working properly.[/yellow]")
                    continue
            except sr.RequestError as e:
                console.print(f"[red]Could not request results from speech recognition service: {e}[/red]")
                console.print("[yellow]Trying offline recognition if available...[/yellow]")
                try:
                    # Try offline recognition if possible
                    if hasattr(recognizer, 'recognize_sphinx'):
                        with sr.Microphone() as source:
                            audio = recognizer.listen(source, timeout=5)
                            speech_text = recognizer.recognize_sphinx(audio)
                            console.print(f"[green]You said (offline recognition):[/green] {speech_text}")
                            user_message = speech_text
                            break
                    else:
                        console.print("[red]Offline recognition not available. Try again later.[/red]")
                        continue
                except Exception:
                    console.print("[red]Offline recognition failed. Please try again.[/red]")
                    continue
            except Exception as e:
                console.print(f"[red]An error occurred: {str(e)}[/red]")
                continue
    elif VOICE_AVAILABLE and user_message.lower() == 'voice output':
        if len(conversation) < 2 or conversation[-1]["role"] != "assistant":
            console.print("[yellow]No AI response available to read.[/yellow]")
            continue
        
        last_response = conversation[-1]["content"]
        console.print("[yellow]Choose voice output method:[/yellow]")
        console.print("1. Basic TTS (default)")
        if CLOUD_TTS_AVAILABLE:
            console.print("2. Cloud TTS (higher quality)")
        choice = input("Select option (or press Enter for default): ")
        
        # Use the enhanced TTS function
        use_cloud = choice == "2" and CLOUD_TTS_AVAILABLE
        
        # Get current settings or use defaults
        rate = getattr(engine, "rate", 150) if "engine" in locals() else 150
        volume = getattr(engine, "volume", 1.0) if "engine" in locals() else 1.0
        
        success = enhanced_tts(last_response, use_cloud=use_cloud, rate=rate, volume=volume)
        
        if not success:
            console.print("[red]Failed to use text-to-speech. Please check your audio settings.[/red]")
        continue
    
    # Add user message to conversation
    conversation.append({"role": "user", "content": user_message})
    
    # Get AI response
    try:
        console.print("\n[bold cyan]AI:[/bold cyan] ", end="")
        
        # Use streaming for more interactive responses
        full_response = ""
        for chunk in client.chat.completions.create(
            messages=conversation,
            model=current_model,
            temperature=temperature,
            stream=True
        ):
            content = chunk.choices[0].delta.content
            if content:
                # Print each chunk as it comes in
                console.print(content, end="", highlight=False)
                full_response += content
        
        console.print()  # Add a new line after response
        
        # Process the full response to find and highlight clickable links
        urls = extract_urls(full_response)
        
        if urls:
            # If we found URLs, show them in a special colorized clickable section
            console.print("\n[bold yellow]Links found in response:[/bold yellow]")
            for i, url in enumerate(urls, 1):
                console.print(f"  [bold blue][link={url}]{i}. {url}[/link][/bold blue]")

        # For better readability, render the response as markdown in a panel
        console.print("\n[bold yellow]Rendered Markdown:[/bold yellow]")
        try:
            md = Markdown(full_response)
            console.print(Panel(md, border_style="green", title="AI Response", subtitle="Markdown Rendering"))
        except Exception as md_error:
            # If markdown rendering fails, just continue without showing an error
            pass
        
        # Add AI response to conversation history
        conversation.append({"role": "assistant", "content": full_response})
        
    except Exception as e:
        console.print(f"\n[bold red]Error:[/bold red] {e}")
        console.print("[yellow]Something went wrong. Please try again.[/yellow]")
# The above code creates a continuous chat with the Groq AI that maintains conversation history

